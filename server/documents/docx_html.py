"""HTML to Word, via python-docx — the engine behind create_docx.

Replaces macOS `textutil`, which converted through the Cocoa HTML importer
and was lossy in ways nothing reported: table cell shading vanished, fonts
resolved to whatever the system had rather than what the CSS asked for, and
point sizes arrived inflated by 4/3 because the importer converted at 96
DPI. Undoing that damage took regex surgery on the emitted OOXML. Building
the document directly means the file says what we meant the first time, and
it works off macOS, so the docx tests finally run on every CI runner instead
of skipping on Linux.

The supported vocabulary is deliberately the one create_docx advertises,
not "HTML":

    headings      h1-h6                (h1-h3 carry the house sizes)
    blocks        p, div, br, hr, pre, blockquote
    inline        b, strong, i, em, u, s/strike/del, code/tt/samp/kbd,
                  sup, sub, a (rendered as its text), span
    lists         ul, ol, li           (nested, to three levels)
    tables        table, tr, td, th, thead, tbody, tfoot, caption

Unknown tags are transparent: their text still renders, so an unexpected
wrapper degrades to plain content instead of vanishing. Anything genuinely
beyond this vocabulary is office_script's job, and create_docx's description
says so.

Cell shading IS honored here (`bgcolor`, or `background-color` in a `style`
attribute), because it costs three lines with python-docx and its silent
loss under textutil is what exposed the whole gap. That is a floor, not the
feature set: office_script remains the answer for merges, borders, images,
and everything else.
"""

import re
from html.parser import HTMLParser

# Tags that close an open paragraph and start their own block.
_BLOCK_TAGS = {
    "p",
    "div",
    "h1",
    "h2",
    "h3",
    "h4",
    "h5",
    "h6",
    "ul",
    "ol",
    "li",
    "table",
    "tr",
    "td",
    "th",
    "blockquote",
    "pre",
    "hr",
    "caption",
}

# Inline tags to the run attribute they switch on.
_INLINE_STYLES = {
    "b": "bold",
    "strong": "bold",
    "i": "italic",
    "em": "italic",
    "u": "underline",
    "s": "strike",
    "strike": "strike",
    "del": "strike",
}

_MONOSPACE_TAGS = {"code", "tt", "samp", "kbd", "pre"}
_MONOSPACE_FONT = "Consolas"

# Deepest list level with a distinct built-in style in the python-docx
# template ("List Bullet 3" exists, "List Bullet 4" does not).
_MAX_LIST_DEPTH = 3

_HEX_RE = re.compile(r"^#?([0-9A-Fa-f]{6})$")
_SHORT_HEX_RE = re.compile(r"^#([0-9A-Fa-f]{3})$")
_RGB_RE = re.compile(r"^rgba?\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)\s*(?:,[^)]*)?\)$")
_BACKGROUND_RE = re.compile(r"background(?:-color)?\s*:\s*([^;]+)", re.I)

# CSS named colors worth recognizing. Not the full list of 148: these are the
# ones a model actually reaches for when asked to color something, and an
# unrecognized name simply leaves the cell unshaded rather than erroring.
_NAMED_COLORS = {
    "black": "000000",
    "white": "FFFFFF",
    "red": "FF0000",
    "green": "008000",
    "lime": "00FF00",
    "blue": "0000FF",
    "yellow": "FFFF00",
    "orange": "FFA500",
    "purple": "800080",
    "gray": "808080",
    "grey": "808080",
    "silver": "C0C0C0",
    "navy": "000080",
    "teal": "008080",
    "olive": "808000",
    "maroon": "800000",
    "aqua": "00FFFF",
    "cyan": "00FFFF",
    "fuchsia": "FF00FF",
    "magenta": "FF00FF",
    "pink": "FFC0CB",
    "brown": "A52A2A",
    "gold": "FFD700",
    "beige": "F5F5DC",
    "ivory": "FFFFF0",
}


def parse_color(raw: str) -> str | None:
    """A CSS color to a 6-digit uppercase hex string, or None if it isn't
    one we recognize. Accepts #rgb, #rrggbb, bare rrggbb (the HTML 3.2
    `bgcolor` spelling), rgb()/rgba(), and the common named colors."""
    if not raw:
        return None
    value = raw.strip().rstrip(";").strip()
    if not value:
        return None
    lowered = value.lower()
    if lowered in _NAMED_COLORS:
        return _NAMED_COLORS[lowered]
    short = _SHORT_HEX_RE.match(value)
    if short:
        r, g, b = short.group(1)
        return (r + r + g + g + b + b).upper()
    exact = _HEX_RE.match(value)
    if exact:
        return exact.group(1).upper()
    rgb = _RGB_RE.match(lowered)
    if rgb:
        channels = [min(255, int(c)) for c in rgb.groups()]
        return "".join(f"{c:02X}" for c in channels)
    return None


def cell_fill(attrs: dict) -> str | None:
    """The shading color for a td/th, from `bgcolor` or a `background-color`
    (or `background`) declaration in `style`. None when neither is present."""
    direct = parse_color(attrs.get("bgcolor", ""))
    if direct:
        return direct
    match = _BACKGROUND_RE.search(attrs.get("style", "") or "")
    return parse_color(match.group(1)) if match else None


def _shade(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


class _Node:
    """Minimal element tree. html.parser is a stream, but tables need to be
    counted (rows x columns) before python-docx can create one, so the
    markup is buffered into a tree first and rendered in a second pass."""

    __slots__ = ("tag", "attrs", "children", "text")

    def __init__(self, tag: str, attrs: dict | None = None, text: str = ""):
        self.tag = tag
        self.attrs = attrs or {}
        self.children: list[_Node] = []
        self.text = text


class _TreeBuilder(HTMLParser):
    """Builds a _Node tree, auto-closing unclosed inline tags at the end.

    Tolerant by design: the input is model-written HTML, and a stray
    unclosed <b> should cost bold formatting on the tail, not the document.
    """

    # Tags that never have a closing partner.
    _VOID = {"br", "hr", "img", "meta", "link", "input", "col", "area"}
    # Tags whose own kind implicitly closes them (<li> after an unclosed
    # <li>), mapped to (what it closes, where the search must stop).
    #
    # The stop set is what makes nesting work. Without it, the <li> opening a
    # SUBLIST would scan past its own <ul> and close the outer <li>, hoisting
    # the sublist to the parent level — the nested list silently flattens.
    _IMPLICIT_CLOSE = {
        "li": ({"li"}, {"ul", "ol"}),
        "td": ({"td", "th"}, {"table"}),
        "th": ({"td", "th"}, {"table"}),
        "tr": ({"tr"}, {"table"}),
    }

    def __init__(self):
        super().__init__(convert_charrefs=True)
        self.root = _Node("#root")
        self._stack = [self.root]

    def handle_starttag(self, tag, attrs):
        tag = tag.lower()
        implicit = self._IMPLICIT_CLOSE.get(tag)
        if implicit:
            closes, stop_at = implicit
            for i in range(len(self._stack) - 1, 0, -1):
                if self._stack[i].tag in stop_at:
                    break
                if self._stack[i].tag in closes:
                    del self._stack[i:]
                    break
        node = _Node(tag, {k.lower(): (v or "") for k, v in attrs})
        self._stack[-1].children.append(node)
        if tag not in self._VOID:
            self._stack.append(node)

    def handle_startendtag(self, tag, attrs):
        tag = tag.lower()
        self._stack[-1].children.append(_Node(tag, {k.lower(): (v or "") for k, v in attrs}))

    def handle_endtag(self, tag):
        tag = tag.lower()
        for i in range(len(self._stack) - 1, 0, -1):
            if self._stack[i].tag == tag:
                del self._stack[i:]
                return
        # Unmatched close tag: ignore it rather than unwinding the document.

    def handle_data(self, data):
        if data:
            self._stack[-1].children.append(_Node("#text", text=data))


class _RunStyle:
    """Inline formatting in effect at some point in the tree."""

    __slots__ = ("bold", "italic", "underline", "strike", "mono", "color")

    def __init__(self, **kw):
        self.bold = kw.get("bold", False)
        self.italic = kw.get("italic", False)
        self.underline = kw.get("underline", False)
        self.strike = kw.get("strike", False)
        self.mono = kw.get("mono", False)
        self.color = kw.get("color")

    def derive(self, **kw) -> "_RunStyle":
        merged = {
            "bold": self.bold,
            "italic": self.italic,
            "underline": self.underline,
            "strike": self.strike,
            "mono": self.mono,
            "color": self.color,
        }
        merged.update(kw)
        return _RunStyle(**merged)


class _Renderer:
    def __init__(self, doc, body_font: str, body_size_pt: float):
        self.doc = doc
        self.body_font = body_font
        self.body_size_pt = body_size_pt
        self._paragraph = None

    # -- run emission ---------------------------------------------------

    def _ensure_paragraph(self, style: str | None = None):
        if self._paragraph is None:
            self._paragraph = self.doc.add_paragraph(style=style)
        return self._paragraph

    def _close_paragraph(self):
        self._paragraph = None

    def _add_text(self, text: str, style: _RunStyle, paragraph=None):
        from docx.shared import Pt, RGBColor

        target = paragraph if paragraph is not None else self._ensure_paragraph()
        run = target.add_run(text)
        run.bold = style.bold or None
        run.italic = style.italic or None
        run.underline = style.underline or None
        if style.strike:
            run.font.strike = True
        run.font.name = _MONOSPACE_FONT if style.mono else self.body_font
        run.font.size = Pt(self.body_size_pt)
        if style.color:
            run.font.color.rgb = RGBColor.from_string(style.color)

    # -- tree walking ---------------------------------------------------

    def render(self, root: _Node) -> None:
        self._walk_children(root, _RunStyle(), container=None)

    def _walk_children(self, node: _Node, style: _RunStyle, container, list_depth: int = 0):
        for child in node.children:
            self._walk(child, style, container, list_depth)

    def _walk(self, node: _Node, style: _RunStyle, container, list_depth: int):
        tag = node.tag

        if tag == "#text":
            text = node.text
            if container is not None:
                # Inside a table cell: whitespace-only text between tags is
                # markup indentation, not content.
                if not text.strip():
                    return
                self._add_text(text, style, paragraph=container)
                return
            if not text.strip() and self._paragraph is None:
                return
            self._add_text(text, style)
            return

        if tag in _INLINE_STYLES:
            self._walk_children(
                node, style.derive(**{_INLINE_STYLES[tag]: True}), container, list_depth
            )
            return

        if tag in _MONOSPACE_TAGS and tag != "pre":
            self._walk_children(node, style.derive(mono=True), container, list_depth)
            return

        if tag in ("span", "a", "font"):
            color = parse_color(node.attrs.get("color", ""))
            if not color:
                match = re.search(r"(?<!-)\bcolor\s*:\s*([^;]+)", node.attrs.get("style", "") or "")
                color = parse_color(match.group(1)) if match else None
            self._walk_children(
                node, style.derive(color=color) if color else style, container, list_depth
            )
            return

        if tag == "br":
            target = container if container is not None else self._ensure_paragraph()
            target.add_run().add_break()
            return

        if tag == "hr":
            self._close_paragraph()
            self.doc.add_paragraph("_" * 40)
            self._close_paragraph()
            return

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6") and container is None:
            self._close_paragraph()
            self._paragraph = self.doc.add_heading("", level=int(tag[1]))
            self._walk_children(node, style, container, list_depth)
            self._close_paragraph()
            return

        if tag in ("ul", "ol"):
            self._close_paragraph()
            self._render_list(node, style, container, list_depth + 1)
            return

        if tag == "table" and container is None:
            self._close_paragraph()
            self._render_table(node, style)
            return

        if tag == "blockquote":
            self._close_paragraph()
            self._paragraph = self.doc.add_paragraph(style="Quote")
            self._walk_children(node, style, container, list_depth)
            self._close_paragraph()
            return

        if tag == "pre":
            self._close_paragraph()
            self._paragraph = self.doc.add_paragraph()
            self._walk_children(node, style.derive(mono=True), container, list_depth)
            self._close_paragraph()
            return

        if tag in _BLOCK_TAGS and container is None:
            self._close_paragraph()
            self._walk_children(node, style, container, list_depth)
            self._close_paragraph()
            return

        # Unknown or already-inside-a-cell element: stay transparent so its
        # text still reaches the document.
        self._walk_children(node, style, container, list_depth)

    def _render_list(self, node: _Node, style: _RunStyle, container, depth: int):
        base = "List Number" if node.tag == "ol" else "List Bullet"
        level = min(depth, _MAX_LIST_DEPTH)
        style_name = base if level == 1 else f"{base} {level}"
        for item in node.children:
            if item.tag != "li":
                # Text or a nested list directly under ul/ol — render it in
                # place rather than dropping it.
                self._walk(item, style, container, depth)
                continue
            self._paragraph = self.doc.add_paragraph(style=style_name)
            for child in item.children:
                if child.tag in ("ul", "ol"):
                    self._close_paragraph()
                    self._render_list(child, style, container, depth + 1)
                else:
                    self._walk(child, style, container, depth)
            self._close_paragraph()

    def _collect_rows(self, node: _Node) -> list[list[_Node]]:
        """Every tr in the table, including those wrapped in thead/tbody/tfoot."""
        rows: list[list[_Node]] = []
        for child in node.children:
            if child.tag == "tr":
                rows.append([c for c in child.children if c.tag in ("td", "th")])
            elif child.tag in ("thead", "tbody", "tfoot"):
                rows.extend(self._collect_rows(child))
        return rows

    def _render_table(self, node: _Node, style: _RunStyle):
        rows = self._collect_rows(node)
        if not rows:
            return
        n_cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=n_cols)
        table.style = "Table Grid"
        for r, cells in enumerate(rows):
            for c, cell_node in enumerate(cells):
                cell = table.cell(r, c)
                paragraph = cell.paragraphs[0]
                cell_style = style.derive(bold=True) if cell_node.tag == "th" else style
                self._walk_children(cell_node, cell_style, container=paragraph)
                fill = cell_fill(cell_node.attrs)
                if fill:
                    _shade(cell, fill)
        # A table and the text after it must not run together.
        self.doc.add_paragraph()
        self._close_paragraph()

        caption = next((c for c in node.children if c.tag == "caption"), None)
        if caption is not None:
            self._paragraph = self.doc.add_paragraph()
            self._walk_children(caption, style.derive(italic=True), container=None)
            self._close_paragraph()


def render_html(doc, html: str, *, body_font: str, body_size_pt: float) -> None:
    """Render `html` into an open python-docx Document, in place."""
    builder = _TreeBuilder()
    builder.feed(html)
    builder.close()
    _Renderer(doc, body_font, body_size_pt).render(builder.root)
