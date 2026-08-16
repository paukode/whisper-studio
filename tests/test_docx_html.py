"""The HTML to Word renderer behind create_docx.

Unit-level companion to test_document_tools.py, which drives the same code
through the executor. What is pinned here is the parsing behaviour that a
model's imperfect HTML actually exercises: color spellings, unclosed tags,
implicit closes, and list nesting. Those are the parts most likely to
regress silently, because a mis-parse produces a perfectly valid document
that is simply missing something.
"""

import pytest
from docx import Document

from server.documents.docx_html import cell_fill, parse_color, render_html


def _render(html: str):
    doc = Document()
    render_html(doc, html, body_font="Calibri", body_size_pt=11)
    return doc


def _styles(doc) -> list[tuple[str, str]]:
    return [(p.style.name, p.text) for p in doc.paragraphs if p.text.strip()]


# ── Colors ─────────────────────────────────────────────────────────────


@pytest.mark.parametrize(
    "raw, expected",
    [
        ("#E63946", "E63946"),
        ("e63946", "E63946"),  # the bare HTML 3.2 bgcolor spelling
        ("#abc", "AABBCC"),
        ("rgb(0, 128, 0)", "008000"),
        ("rgba(255,255,0,0.5)", "FFFF00"),
        ("rgb(300,0,0)", "FF0000"),  # clamped, not wrapped
        ("Yellow", "FFFF00"),
        ("  #ffffff ; ", "FFFFFF"),
        ("", None),
        ("not-a-color", None),
        ("#12345", None),
    ],
)
def test_parse_color(raw, expected):
    assert parse_color(raw) == expected


@pytest.mark.parametrize(
    "attrs, expected",
    [
        ({"bgcolor": "#123456"}, "123456"),
        ({"style": "background-color:#123456"}, "123456"),
        ({"style": "color:red;background:blue"}, "0000FF"),
        # bgcolor wins when both are present; neither present means no fill.
        ({"bgcolor": "red", "style": "background:blue"}, "FF0000"),
        ({"style": "color:red"}, None),
        ({}, None),
    ],
)
def test_cell_fill(attrs, expected):
    assert cell_fill(attrs) == expected


def test_text_color_applies_to_runs():
    doc = _render('<p>plain <span style="color:#FF0000">red</span></p>')
    runs = {r.text: r for r in doc.paragraphs[0].runs}
    assert runs["red"].font.color.rgb is not None
    assert str(runs["red"].font.color.rgb) == "FF0000"
    assert runs["plain "].font.color.rgb is None


# ── Structure ──────────────────────────────────────────────────────────


def test_nested_lists_step_down_a_level():
    doc = _render("<ul><li>A<ul><li>B<ul><li>C</li></ul></li></ul></li></ul>")
    assert _styles(doc) == [
        ("List Bullet", "A"),
        ("List Bullet 2", "B"),
        ("List Bullet 3", "C"),
    ]


def test_list_nesting_stops_at_the_deepest_template_style():
    """Word's default template defines List Bullet 1-3 only; a fourth level
    must reuse the third rather than raise on a missing style."""
    doc = _render("<ul><li>1<ul><li>2<ul><li>3<ul><li>4</li></ul></li></ul></li></ul></li></ul>")
    assert [s for s, _ in _styles(doc)][-1] == "List Bullet 3"


def test_unclosed_list_items_still_separate():
    assert _styles(_render("<ul><li>A<li>B<li>C</ul>")) == [
        ("List Bullet", "A"),
        ("List Bullet", "B"),
        ("List Bullet", "C"),
    ]


def test_ordered_list_inside_unordered_keeps_both_kinds():
    doc = _render("<ul><li>Outer<ol><li>One</li></ol></li></ul>")
    assert _styles(doc) == [("List Bullet", "Outer"), ("List Number 2", "One")]


def test_table_rows_are_collected_through_thead_and_tbody():
    doc = _render(
        "<table><thead><tr><th>H1</th><th>H2</th></tr></thead>"
        "<tbody><tr><td>a</td><td>b</td></tr></tbody></table>"
    )
    table = doc.tables[0]
    assert (len(table.rows), len(table.columns)) == (2, 2)
    assert [c.text for c in table.rows[1].cells] == ["a", "b"]


def test_ragged_table_pads_to_the_widest_row():
    doc = _render("<table><tr><td>a</td></tr><tr><td>b</td><td>c</td></tr></table>")
    table = doc.tables[0]
    assert len(table.columns) == 2
    assert table.rows[0].cells[1].text == ""


def test_nested_table_does_not_swallow_the_outer_row():
    """A <td> opening inside a nested table must not implicitly close the
    outer table's cell — the same scoping bug that flattens nested lists."""
    doc = _render(
        "<table><tr><td>outer<table><tr><td>inner</td></tr></table></td>"
        "<td>sibling</td></tr></table>"
    )
    outer = doc.tables[0]
    assert len(outer.rows[0].cells) == 2
    assert outer.rows[0].cells[1].text == "sibling"


def test_empty_table_is_skipped_not_crashed():
    assert _render("<table></table>").tables == []


def test_unknown_tags_stay_transparent():
    doc = _render("<article><section>kept</section></article>")
    assert "kept" in " ".join(p.text for p in doc.paragraphs)


def test_entities_are_decoded():
    doc = _render("<p>a &amp; b &lt;c&gt; &#233;</p>")
    assert doc.paragraphs[0].text == "a & b <c> é"


def test_br_becomes_a_line_break_not_a_paragraph():
    doc = _render("<p>one<br>two</p>")
    assert len([p for p in doc.paragraphs if p.text.strip()]) == 1
    assert "one" in doc.paragraphs[0].text and "two" in doc.paragraphs[0].text


def test_stray_close_tag_does_not_unwind_the_document():
    doc = _render("</div></p><p>survives</p>")
    assert "survives" in " ".join(p.text for p in doc.paragraphs)


def test_markup_whitespace_between_cells_is_not_content():
    doc = _render("<table>\n  <tr>\n    <td>a</td>\n  </tr>\n</table>")
    assert doc.tables[0].rows[0].cells[0].text == "a"
